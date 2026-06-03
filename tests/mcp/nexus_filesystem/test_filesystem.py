"""Adversarial tests for nexus-filesystem MCP server (Phase 2.5a Server 1)."""
from __future__ import annotations

import io
import os
from pathlib import Path

import pytest

from mcp_servers.nexus_filesystem import (
    append_to_file,
    delete_file,
    file_info,
    find_file,
    list_directory,
    move_file,
    read_docx,
    read_image_metadata,
    read_pdf,
    read_spreadsheet,
    read_text_file,
    recent_files,
    view_image,
    write_text_file,
)


SANDBOX_PREFIX = "_mcp_fs_test_"


def _allow_tmp_workspace(monkeypatch, tmp_path: Path) -> Path:
    workspace = tmp_path.resolve()
    monkeypatch.setattr("services.file_system_tools.ALLOWED_ROOTS", (workspace,))
    monkeypatch.setattr("mcp_servers.nexus_filesystem.ALLOWED_ROOTS", (workspace,))
    return workspace


@pytest.fixture
def sandbox(tmp_path, monkeypatch):
    """Provide a sandboxed directory inside ALLOWED_ROOTS for write tests."""
    workspace = _allow_tmp_workspace(monkeypatch, tmp_path)
    d = workspace / SANDBOX_PREFIX.rstrip("_")
    d.mkdir()
    return d


# ---------- read_text_file ----------

def test_read_text_file_outside_root_rejected():
    r = read_text_file(path="/etc/passwd")
    assert r["ok"] is False
    assert r["reason"] == "denied"


def test_read_text_file_relative_rejected():
    r = read_text_file(path="relative.txt")
    assert r["ok"] is False


def test_read_text_file_happy_path(sandbox):
    target = sandbox / "hello.txt"
    target.write_text("hello mcp world", encoding="utf-8")
    r = read_text_file(path=str(target))
    assert r["ok"] is True
    assert r["content"] == "hello mcp world"
    assert r["kind"] == "text"


# ---------- read_pdf / read_docx / read_spreadsheet ----------

def test_read_pdf_rejects_non_pdf_extension(sandbox):
    target = sandbox / "not_a_pdf.txt"
    target.write_text("hi")
    r = read_pdf(path=str(target))
    assert r["ok"] is False
    assert r["reason"] == "not_a_pdf"


def test_read_pdf_path_outside_root_rejected():
    r = read_pdf(path="/etc/something.pdf")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


def test_read_pdf_happy_path(sandbox):
    """Generate a tiny PDF with reportlab and read it back."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not yet installed (added in pdf-docs server)")
    target = sandbox / "demo.pdf"
    c = canvas.Canvas(str(target))
    c.drawString(72, 720, "MCP test PDF content")
    c.showPage()
    c.save()
    r = read_pdf(path=str(target))
    assert r["ok"] is True
    assert "MCP test PDF content" in r["text"]
    assert r["page_count"] >= 1


def test_read_docx_rejects_non_docx(sandbox):
    target = sandbox / "x.txt"
    target.write_text("y")
    r = read_docx(path=str(target))
    assert r["ok"] is False
    assert r["reason"] == "not_a_docx"


def test_read_docx_happy_path(sandbox):
    from docx import Document
    target = sandbox / "demo.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(target))
    r = read_docx(path=str(target))
    assert r["ok"] is True
    assert r["paragraph_count"] == 2
    assert "First paragraph" in r["text"]


def test_read_spreadsheet_csv_happy(sandbox):
    target = sandbox / "data.csv"
    target.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
    r = read_spreadsheet(path=str(target))
    assert r["ok"] is True
    assert r["kind"] == "csv"
    assert r["header"] == ["name", "age"]
    assert r["rows"] == [["Alice", "30"], ["Bob", "25"]]


def test_read_spreadsheet_xlsx_happy(sandbox):
    import openpyxl
    target = sandbox / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["product", "qty"])
    ws.append(["wallet", 5])
    ws.append(["bag", 2])
    wb.save(str(target))
    r = read_spreadsheet(path=str(target))
    assert r["ok"] is True
    assert r["kind"] == "xlsx"
    assert r["header"] == ["product", "qty"]
    assert r["row_count"] == 2


def test_read_spreadsheet_unsupported_extension(sandbox):
    target = sandbox / "data.txt"
    target.write_text("hi")
    r = read_spreadsheet(path=str(target))
    assert r["ok"] is False
    assert r["reason"] == "unsupported_extension"


def test_read_spreadsheet_outside_root_rejected():
    r = read_spreadsheet(path="/etc/spreadsheet.csv")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


# ---------- read_image_metadata / view_image ----------

def _make_png(target: Path) -> None:
    from PIL import Image
    img = Image.new("RGB", (8, 8), color=(100, 200, 50))
    img.save(target, "PNG")


def test_read_image_metadata_happy(sandbox):
    target = sandbox / "test.png"
    _make_png(target)
    r = read_image_metadata(path=str(target))
    assert r["ok"] is True
    assert r["format"] == "PNG"
    assert r["width"] == 8
    assert r["height"] == 8


def test_read_image_metadata_outside_root_rejected():
    r = read_image_metadata(path="/etc/some.png")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


def test_view_image_happy(sandbox):
    target = sandbox / "vi.png"
    _make_png(target)
    r = view_image(path=str(target))
    assert r["ok"] is True
    assert r["mime_type"] == "image/png"
    assert isinstance(r["base64"], str) and len(r["base64"]) > 0


def test_view_image_outside_root_rejected():
    r = view_image(path="/etc/image.png")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


# ---------- list_directory ----------

def test_list_directory_basic():
    r = list_directory(path=str(Path.home()))
    assert r["ok"] is True
    assert isinstance(r["entries"], list)


def test_list_directory_denied():
    r = list_directory(path="/etc")
    assert r["ok"] is False
    assert r["reason"] == "denied"


# ---------- find_file ----------

def test_find_file_empty_query_rejected():
    r = find_file(query="")
    assert r["ok"] is False
    assert r["reason"] == "empty_query"


def test_find_file_outside_root_rejected():
    r = find_file(query="x", search_root="/etc")
    assert r["ok"] is False


def test_find_file_by_name_happy(sandbox):
    (sandbox / "needle_marker.txt").write_text("x")
    (sandbox / "haystack_other.txt").write_text("y")
    r = find_file(query="needle_marker", search_root=str(sandbox))
    assert r["ok"] is True
    assert r["match_count"] >= 1
    assert any("needle_marker" in p for p in r["matches"])


def test_find_file_content_search(sandbox):
    (sandbox / "a.txt").write_text("UNIQUE_TOKEN_PHASE_2_5_A inside this file")
    (sandbox / "b.txt").write_text("nothing here")
    r = find_file(query="UNIQUE_TOKEN_PHASE_2_5_A", search_root=str(sandbox), content=True)
    assert r["ok"] is True
    assert any("a.txt" in h["path"] for h in r["hits"])
    assert not any("b.txt" in h["path"] for h in r["hits"])


# ---------- recent_files ----------

def test_recent_files_days_out_of_range_rejected():
    r = recent_files(days=0)
    assert r["ok"] is False


def test_recent_files_outside_root_rejected():
    r = recent_files(search_root="/etc", days=1)
    assert r["ok"] is False


def test_recent_files_includes_just_written(sandbox):
    target = sandbox / "fresh.md"
    target.write_text("just made")
    r = recent_files(extension=".md", days=1, search_root=str(sandbox))
    assert r["ok"] is True
    assert any("fresh.md" in res["path"] for res in r["results"])


def test_recent_files_extension_filter_excludes(sandbox):
    (sandbox / "fresh.py").write_text("just made")
    r = recent_files(extension=".md", days=1, search_root=str(sandbox))
    assert r["ok"] is True
    assert not any(res["path"].endswith(".py") for res in r["results"])


# ---------- file_info ----------

def test_file_info_happy(sandbox):
    target = sandbox / "info.txt"
    target.write_text("hello")
    r = file_info(path=str(target))
    assert r["ok"] is True
    assert r["is_file"] is True
    assert r["size_bytes"] == 5


def test_file_info_outside_root_rejected():
    r = file_info(path="/etc/passwd")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


# ---------- write_text_file ----------

def test_write_text_file_outside_root_rejected():
    r = write_text_file(path="/etc/x.txt", content="hello")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


def test_write_text_file_relative_rejected():
    r = write_text_file(path="x.txt", content="hello")
    assert r["ok"] is False


def test_write_text_file_denied_filename_rejected(sandbox):
    r = write_text_file(path=str(sandbox / ".env"), content="SECRET=1")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


def test_write_text_file_too_large_rejected(sandbox):
    big = "x" * (260 * 1024)  # > 256 KB cap
    r = write_text_file(path=str(sandbox / "big.txt"), content=big)
    assert r["ok"] is False
    assert r["reason"] == "content_too_large"


def test_write_text_file_happy(sandbox):
    target = sandbox / "out.txt"
    r = write_text_file(path=str(target), content="written by MCP")
    assert r["ok"] is True
    assert target.read_text(encoding="utf-8") == "written by MCP"


def test_write_text_file_non_string_content_rejected(sandbox):
    r = write_text_file(path=str(sandbox / "x.txt"), content=123)
    assert r["ok"] is False


# ---------- append_to_file ----------

def test_append_to_file_happy(sandbox):
    target = sandbox / "log.txt"
    write_text_file(path=str(target), content="line1\n")
    r = append_to_file(path=str(target), content="line2\n")
    assert r["ok"] is True
    assert target.read_text() == "line1\nline2\n"


def test_append_to_file_outside_root_rejected():
    r = append_to_file(path="/etc/x.txt", content="hi")
    assert r["ok"] is False


# ---------- move_file ----------

def test_move_file_happy(sandbox):
    src = sandbox / "a.txt"
    src.write_text("contents")
    dst = sandbox / "b.txt"
    r = move_file(from_path=str(src), to_path=str(dst))
    assert r["ok"] is True
    assert not src.exists()
    assert dst.exists()


def test_move_file_destination_exists_rejected(sandbox):
    src = sandbox / "src.txt"
    src.write_text("a")
    dst = sandbox / "dst.txt"
    dst.write_text("existing")
    r = move_file(from_path=str(src), to_path=str(dst))
    assert r["ok"] is False
    assert r["reason"] == "destination_exists"


def test_move_file_source_outside_root_rejected(sandbox):
    r = move_file(from_path="/etc/hosts", to_path=str(sandbox / "h"))
    assert r["ok"] is False
    assert r["reason"] == "from_path_denied"


def test_move_file_destination_outside_root_rejected(sandbox):
    src = sandbox / "a.txt"
    src.write_text("a")
    r = move_file(from_path=str(src), to_path="/etc/dest.txt")
    assert r["ok"] is False
    assert r["reason"] == "to_path_denied"


# ---------- delete_file ----------

def test_delete_file_happy(sandbox):
    target = sandbox / "doomed.txt"
    target.write_text("bye")
    r = delete_file(path=str(target))
    assert r["ok"] is True
    assert not target.exists()


def test_delete_file_outside_root_rejected():
    r = delete_file(path="/etc/passwd")
    assert r["ok"] is False
    assert r["reason"] == "path_denied"


def test_delete_file_refuses_directory(sandbox):
    sub = sandbox / "subdir"
    sub.mkdir()
    r = delete_file(path=str(sub))
    assert r["ok"] is False
    assert r["reason"] == "is_directory"


def test_delete_file_denied_filename(sandbox):
    # .pem files are denied by Phase 4 pattern — validation fails first
    target = sandbox / "fake.pem"
    target.write_text("not really")
    r = delete_file(path=str(target))
    assert r["ok"] is False
    # Either validation rejects the path entirely or filename pattern denies the delete.
    assert r["reason"] in {"path_denied", "filename_denied"}
